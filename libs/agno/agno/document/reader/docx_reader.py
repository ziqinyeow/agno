import asyncio
from pathlib import Path
from typing import IO, Any, List, Union

from agno.document.base import Document
from agno.document.reader.base import Reader
from agno.utils.log import log_info, logger

try:
    from docx import Document as DocxDocument  # type: ignore
except ImportError:
    raise ImportError("The `python-docx` package is not installed. Please install it via `pip install python-docx`.")


class DocxReader(Reader):
    """Reader for Doc/Docx files"""

    def read(self, file: Union[Path, IO[Any]]) -> List[Document]:
        """Read a docx file and return a list of documents"""
        try:
            if isinstance(file, Path):
                if not file.exists():
                    raise FileNotFoundError(f"Could not find file: {file}")
                log_info(f"Reading: {file}")
                docx_document = DocxDocument(str(file))
                doc_name = file.stem
            else:
                log_info(f"Reading uploaded file: {file.name}")
                docx_document = DocxDocument(file)
                doc_name = file.name.split(".")[0]

            doc_content = "\n\n".join([para.text for para in docx_document.paragraphs])

            documents = [
                Document(
                    name=doc_name,
                    id=doc_name,
                    content=doc_content,
                )
            ]

            if self.chunk:
                chunked_documents = []
                for document in documents:
                    chunked_documents.extend(self.chunk_document(document))
                return chunked_documents
            return documents

        except Exception as e:
            logger.error(f"Error reading file: {e}")
            return []

    async def async_read(self, file: Union[Path, IO[Any]]) -> List[Document]:
        """Asynchronously read a docx file and return a list of documents"""
        try:
            return await asyncio.to_thread(self.read, file)
        except Exception as e:
            logger.error(f"Error reading file asynchronously: {e}")
            return []
